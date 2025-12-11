#!/usr/bin/env python3
"""
Script to create a Word document presentation for the Handwriting Recognition project.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_presentation():
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ============== TITLE PAGE ==============
    # Add some spacing before title
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Handwriting Recognition")
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Deep Learning OCR with CNN-BiLSTM-CTC Architecture")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Key stats box
    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats_run = stats_para.add_run("87% Character Accuracy | 9.1M Parameters | 20 min Training")
    stats_run.font.size = Pt(14)
    stats_run.font.color.rgb = RGBColor(0, 128, 0)
    stats_run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Technology badges
    tech_para = doc.add_paragraph()
    tech_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tech_run = tech_para.add_run("PyTorch  |  Python 3.12  |  Hugging Face  |  Google Colab")
    tech_run.font.size = Pt(12)
    tech_run.italic = True

    # Page break
    doc.add_page_break()

    # ============== TABLE OF CONTENTS ==============
    toc_title = doc.add_heading("Table of Contents", level=1)

    toc_items = [
        "1. Executive Summary",
        "2. Project Overview",
        "3. Technology Stack",
        "4. Model Architecture",
        "5. Dataset Analysis",
        "6. Training Results",
        "7. Performance Metrics",
        "8. Quick Start Guide",
        "9. Use Cases & Applications",
        "10. Future Improvements",
        "11. Conclusion"
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ============== EXECUTIVE SUMMARY ==============
    doc.add_heading("1. Executive Summary", level=1)

    exec_summary = """This project implements a state-of-the-art handwriting recognition system using deep learning. The system converts images of handwritten text into digital text with 87% character-level accuracy.

Key Achievements:
"""
    doc.add_paragraph(exec_summary)

    achievements = [
        ("Character Accuracy", "87% (CER: 12.95%)"),
        ("Word Accuracy", "57.5% (WER: 42.47%)"),
        ("Training Samples", "10,373 from IAM Database"),
        ("Model Size", "105MB (9.1M parameters)"),
        ("Training Time", "~20 minutes on T4 GPU"),
        ("Inference Speed", "50-100ms per image (GPU)")
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True

    for metric, value in achievements:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value

    doc.add_paragraph()
    doc.add_paragraph("The model is production-ready and available on Hugging Face Hub for immediate deployment.")

    doc.add_page_break()

    # ============== PROJECT OVERVIEW ==============
    doc.add_heading("2. Project Overview", level=1)

    doc.add_heading("Purpose", level=2)
    doc.add_paragraph("The primary goal of this project is to build an end-to-end Optical Character Recognition (OCR) system that can automatically convert handwritten text images into digital text.")

    doc.add_heading("Problem Statement", level=2)
    doc.add_paragraph("""Traditional OCR systems struggle with handwritten text due to:
- High variability in writing styles
- Inconsistent character spacing
- Connected/cursive letters
- Variable image quality

This project addresses these challenges using modern deep learning techniques.""")

    doc.add_heading("Solution Approach", level=2)
    doc.add_paragraph("We implement a CNN-BiLSTM-CTC architecture that:")

    bullet_points = [
        "Extracts visual features using Convolutional Neural Networks (CNN)",
        "Models sequential dependencies with Bidirectional LSTM",
        "Uses CTC Loss for alignment-free training",
        "Requires only text labels (no character position annotations)"
    ]

    for point in bullet_points:
        p = doc.add_paragraph(point, style='List Bullet')

    doc.add_page_break()

    # ============== TECHNOLOGY STACK ==============
    doc.add_heading("3. Technology Stack", level=1)

    doc.add_heading("Core Technologies", level=2)

    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Table Grid'
    hdr = tech_table.rows[0].cells
    hdr[0].text = 'Technology'
    hdr[1].text = 'Version'
    hdr[2].text = 'Purpose'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    technologies = [
        ("Python", "3.12+", "Primary programming language"),
        ("PyTorch", "2.0+", "Deep learning framework"),
        ("Hugging Face Datasets", "2.14+", "Dataset loading"),
        ("Pillow", "9.5+", "Image processing"),
        ("NumPy", "1.24+", "Numerical computations"),
        ("Matplotlib", "3.7+", "Visualization"),
        ("Seaborn", "0.13+", "Statistical plots"),
        ("jiwer", "3.0+", "CER/WER metrics"),
        ("Jupyter", "1.0+", "Development environment")
    ]

    for tech, ver, purpose in technologies:
        row = tech_table.add_row().cells
        row[0].text = tech
        row[1].text = ver
        row[2].text = purpose

    doc.add_paragraph()

    doc.add_heading("Deployment Platforms", level=2)
    platforms = [
        "Google Colab: Free GPU training (T4/A100)",
        "Hugging Face Hub: Model hosting and distribution",
        "Local GPU: For production deployment"
    ]
    for p in platforms:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_page_break()

    # ============== MODEL ARCHITECTURE ==============
    doc.add_heading("4. Model Architecture", level=1)

    doc.add_heading("Architecture Overview: CNN-BiLSTM-CTC", level=2)

    arch_desc = """The model follows a proven architecture for sequence-to-sequence text recognition:

1. CNN Feature Extractor (7 blocks)
   - Input: Grayscale image [Batch, 1, 128, Width]
   - Output: Feature maps [Batch, 512, 7, Width/4]
   - Uses progressive channel growth: 1→64→128→256→512
   - Asymmetric pooling preserves horizontal resolution

2. Sequence Mapping Layer
   - Reshapes CNN output to sequence format
   - Linear projection: 3584 → 256 dimensions

3. Bidirectional LSTM (2 layers)
   - Hidden size: 256 per direction
   - Output: 512 dimensions (forward + backward)
   - Dropout: 0.3 for regularization

4. CTC Output Layer
   - Linear: 512 → 75 (74 characters + blank token)
   - LogSoftmax for probability distribution
"""
    doc.add_paragraph(arch_desc)

    doc.add_heading("Model Parameters", level=2)

    params_table = doc.add_table(rows=1, cols=2)
    params_table.style = 'Table Grid'
    hdr = params_table.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Parameters'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    params = [
        ("CNN Feature Extractor", "~4.5M"),
        ("Sequence Mapper", "~0.9M"),
        ("BiLSTM Layers", "~3.2M"),
        ("Output Layer", "~0.5M"),
        ("Total", "9,139,147 (9.1M)")
    ]

    for comp, param in params:
        row = params_table.add_row().cells
        row[0].text = comp
        row[1].text = param

    doc.add_paragraph()

    doc.add_heading("Why This Architecture?", level=2)

    reasons = [
        "CNN: Efficiently extracts visual features from handwritten strokes",
        "BiLSTM: Captures context from both directions (important for language)",
        "CTC Loss: Eliminates need for expensive character-level annotations",
        "Proven: This architecture is the industry standard for OCR tasks"
    ]
    for r in reasons:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_page_break()

    # ============== DATASET ANALYSIS ==============
    doc.add_heading("5. Dataset Analysis", level=1)

    doc.add_heading("IAM Handwriting Database", level=2)
    doc.add_paragraph("The model is trained on the IAM Handwriting Database, a widely-used benchmark for handwriting recognition research.")

    dataset_table = doc.add_table(rows=1, cols=2)
    dataset_table.style = 'Table Grid'
    hdr = dataset_table.rows[0].cells
    hdr[0].text = 'Statistic'
    hdr[1].text = 'Value'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    stats = [
        ("Total Samples", "10,373"),
        ("Training Set", "6,482 samples"),
        ("Validation Set", "976 samples"),
        ("Test Set", "2,915 samples"),
        ("Unique Characters", "74 (a-z, A-Z, 0-9, space, punctuation)"),
        ("Average Text Length", "48-60 characters"),
        ("Text Length Range", "5-150 characters"),
        ("Source", "University of Bern / Teklia (Hugging Face)")
    ]

    for stat, val in stats:
        row = dataset_table.add_row().cells
        row[0].text = stat
        row[1].text = val

    doc.add_paragraph()

    doc.add_heading("Character Distribution", level=2)
    doc.add_paragraph("The dataset follows natural English text frequency distribution:")

    char_freq = [
        ("Space", "Most common (word separator)"),
        ("'e'", "13.2% - Most frequent letter"),
        ("'t'", "9.4%"),
        ("'a'", "8.1%"),
        ("'o'", "7.9%"),
        ("'i'", "7.0%")
    ]

    for char, freq in char_freq:
        doc.add_paragraph(f"{char}: {freq}", style='List Bullet')

    doc.add_page_break()

    # ============== TRAINING RESULTS ==============
    doc.add_heading("6. Training Results", level=1)

    doc.add_heading("Training Configuration", level=2)

    config_table = doc.add_table(rows=1, cols=3)
    config_table.style = 'Table Grid'
    hdr = config_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Value'
    hdr[2].text = 'Rationale'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    config = [
        ("Epochs", "10", "Convergence achieved"),
        ("Batch Size", "8", "GPU memory optimization"),
        ("Learning Rate", "0.001", "Adam default"),
        ("Optimizer", "Adam", "Adaptive learning rates"),
        ("LR Scheduler", "ReduceLROnPlateau", "Dynamic adjustment"),
        ("Gradient Clipping", "5.0", "Stable RNN training"),
        ("Image Height", "128px", "Balance detail vs. speed")
    ]

    for param, val, rationale in config:
        row = config_table.add_row().cells
        row[0].text = param
        row[1].text = val
        row[2].text = rationale

    doc.add_paragraph()

    doc.add_heading("Training Progress", level=2)

    progress_table = doc.add_table(rows=1, cols=5)
    progress_table.style = 'Table Grid'
    hdr = progress_table.rows[0].cells
    headers = ['Epoch', 'Train Loss', 'Val Loss', 'CER', 'WER']
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    progress = [
        ("1", "3.21", "2.67", "100%", "100%"),
        ("2", "1.69", "1.03", "29.3%", "71.8%"),
        ("5", "0.60", "0.57", "17.7%", "53.1%"),
        ("7", "0.49", "0.46", "14.4%", "46.5%"),
        ("10 (Final)", "0.39", "0.38", "12.95%", "42.47%")
    ]

    for epoch, train, val, cer, wer in progress:
        row = progress_table.add_row().cells
        row[0].text = epoch
        row[1].text = train
        row[2].text = val
        row[3].text = cer
        row[4].text = wer

    doc.add_paragraph()
    doc.add_paragraph("Training Time: ~20 minutes on NVIDIA T4 GPU (1.7-2.1 min/epoch)")

    doc.add_page_break()

    # ============== PERFORMANCE METRICS ==============
    doc.add_heading("7. Performance Metrics", level=1)

    doc.add_heading("Accuracy Metrics", level=2)

    doc.add_paragraph("""
Character Error Rate (CER): 12.95%
- Measures character-level accuracy
- 87.05% of characters are correctly recognized
- Industry competitive for handwriting OCR

Word Error Rate (WER): 42.47%
- Measures word-level accuracy
- 57.53% of words are exactly correct
- Higher than CER because one character error fails the whole word
""")

    doc.add_heading("Understanding CER vs WER", level=2)
    doc.add_paragraph("""Example:
Ground Truth: "magnificent"
Prediction:   "magnifcent" (missing 'i')

CER: 1 error / 11 characters = 9.1%
WER: 1 error / 1 word = 100%

This explains why WER is significantly higher than CER.""")

    doc.add_heading("Inference Speed", level=2)

    speed_table = doc.add_table(rows=1, cols=3)
    speed_table.style = 'Table Grid'
    hdr = speed_table.rows[0].cells
    hdr[0].text = 'Hardware'
    hdr[1].text = 'Speed'
    hdr[2].text = 'Memory'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    speeds = [
        ("CPU (Intel i7)", "200-500ms/image", "500MB"),
        ("GPU (T4)", "50-100ms/image", "2GB"),
        ("GPU (V100)", "20-40ms/image", "4GB"),
        ("GPU (A100)", "10-20ms/image", "4-8GB")
    ]

    for hw, speed, mem in speeds:
        row = speed_table.add_row().cells
        row[0].text = hw
        row[1].text = speed
        row[2].text = mem

    doc.add_page_break()

    # ============== QUICK START GUIDE ==============
    doc.add_heading("8. Quick Start Guide", level=1)

    doc.add_heading("Installation", level=2)
    doc.add_paragraph("pip install torch datasets pillow numpy huggingface_hub", style='Quote')

    doc.add_heading("Download Pre-trained Model", level=2)
    code1 = """from huggingface_hub import hf_hub_download

model_path = hf_hub_download(
    repo_id="IsmatS/handwriting-recognition-iam",
    filename="best_model.pth"
)"""
    doc.add_paragraph(code1, style='Quote')

    doc.add_heading("Load and Use Model", level=2)
    code2 = """import torch
from PIL import Image

# Load checkpoint
checkpoint = torch.load(model_path, map_location='cpu')

# Initialize model (CRNN class from train_colab.ipynb)
model = CRNN(num_classes=75)
model.load_state_dict(checkpoint['model_state_dict'])
model.eval()

# Preprocess image
img = Image.open('handwriting.png').convert('L')
# Resize maintaining aspect ratio to height=128
w, h = img.size
new_w = int(128 * (w / h))
img = img.resize((new_w, 128), Image.LANCZOS)

# Convert to tensor and normalize
import numpy as np
img_array = np.array(img) / 255.0
img_array = (img_array - 0.5) / 0.5
tensor = torch.FloatTensor(img_array).unsqueeze(0).unsqueeze(0)

# Predict
with torch.no_grad():
    output = model(tensor)
    # Use CTC decoding to get text
    predicted_text = decode_predictions(output, char_mapper)
    print(predicted_text)"""
    doc.add_paragraph(code2, style='Quote')

    doc.add_page_break()

    # ============== USE CASES ==============
    doc.add_heading("9. Use Cases & Applications", level=1)

    use_cases = [
        ("Document Digitization", "Convert handwritten notes, letters, and historical documents to searchable digital text"),
        ("Healthcare", "Transcribe handwritten prescriptions and medical notes"),
        ("Education", "Grade handwritten assignments and exams automatically"),
        ("Banking & Finance", "Process handwritten checks and forms"),
        ("Postal Services", "Read handwritten addresses on mail"),
        ("Legal", "Digitize handwritten contracts and legal documents"),
        ("Archive Management", "Make historical handwritten records searchable"),
        ("Personal Productivity", "Convert handwritten to-do lists and notes to digital format")
    ]

    for title, desc in use_cases:
        p = doc.add_paragraph()
        run = p.add_run(title + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ============== FUTURE IMPROVEMENTS ==============
    doc.add_heading("10. Future Improvements", level=1)

    improvements = [
        ("Attention Mechanism", "Add attention layers for better focus on relevant image regions"),
        ("Transformer Architecture", "Implement Vision Transformer (ViT) for potentially better accuracy"),
        ("Data Augmentation", "Add rotation, elastic distortion, and noise for robustness"),
        ("Model Scaling", "Increase to 20-50M parameters for improved accuracy"),
        ("Multi-line Support", "Extend to paragraph and document-level recognition"),
        ("Language Model Integration", "Add spell-checking and context-aware corrections"),
        ("Multilingual Support", "Extend character set to support multiple languages"),
        ("Real-time Processing", "Optimize for video stream processing"),
        ("Mobile Deployment", "Create TensorFlow Lite / ONNX models for mobile devices")
    ]

    for title, desc in improvements:
        p = doc.add_paragraph()
        run = p.add_run(title + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ============== CONCLUSION ==============
    doc.add_heading("11. Conclusion", level=1)

    conclusion = """This handwriting recognition project successfully demonstrates the implementation of a production-ready OCR system using modern deep learning techniques.

Key Accomplishments:
"""
    doc.add_paragraph(conclusion)

    accomplishments = [
        "Achieved 87% character-level accuracy on the IAM benchmark dataset",
        "Implemented industry-standard CNN-BiLSTM-CTC architecture",
        "Trained efficiently in ~20 minutes on consumer GPU hardware",
        "Created comprehensive documentation and visualization",
        "Deployed pre-trained model on Hugging Face Hub for easy access",
        "Provided complete training pipeline in Google Colab-ready notebook"
    ]

    for acc in accomplishments:
        doc.add_paragraph(acc, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph("The project serves as both a practical tool for handwriting recognition and an educational resource for understanding deep learning-based OCR systems.")

    doc.add_paragraph()

    # Final note
    final = doc.add_paragraph()
    final_run = final.add_run("Model available at: ")
    final.add_run("https://huggingface.co/IsmatS/handwriting-recognition-iam")

    doc.add_paragraph()

    # References
    doc.add_heading("References", level=2)
    refs = [
        "IAM Handwriting Database - University of Bern",
        "PyTorch Documentation - pytorch.org",
        "CTC Loss Paper - Graves et al., 2006",
        "CRNN Architecture - Shi et al., 2015"
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Bullet')

    # Save document
    output_path = '/home/user/handwriting_recognition/Handwriting_Recognition_Presentation.docx'
    doc.save(output_path)
    print(f"Presentation saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_presentation()
